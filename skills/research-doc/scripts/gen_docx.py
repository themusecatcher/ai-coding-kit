#!/usr/bin/env python3
"""
Word 调研报告生成器
用法：python3 gen_docx.py --input content.json --output report.docx
"""

import argparse
import json
import sys
import os
from datetime import datetime

def check_deps():
    try:
        import docx
    except ImportError:
        print("❌ 缺少 python-docx，请执行: pip3 install --user python-docx")
        sys.exit(1)

def create_docx(data: dict, output_path: str):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── 样式配置 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3

    # ── 标题页 ──
    title = doc.add_heading(data.get("title", "调研报告"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    if data.get("subtitle"):
        sub = doc.add_paragraph(data["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 元信息
    meta_parts = []
    if data.get("author"):
        meta_parts.append(f"作者：{data['author']}")
    meta_parts.append(f"日期：{data.get('date', datetime.now().strftime('%Y-%m-%d'))}")
    meta = doc.add_paragraph(" | ".join(meta_parts))
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── 目录页（章节 >= 3 时） ──
    sections = data.get("sections", [])
    if len(sections) >= 3:
        doc.add_heading("目录", level=1)
        for i, sec in enumerate(sections):
            toc_item = doc.add_paragraph(f"{i + 1}. {sec.get('title', f'章节 {i + 1}')}")
            toc_item.paragraph_format.space_after = Pt(4)
            for run in toc_item.runs:
                run.font.size = Pt(12)
        doc.add_page_break()

    # ── 内容章节 ──
    for i, section in enumerate(sections):
        sec_type = section.get("type", "text")
        sec_title = section.get("title", f"章节 {i + 1}")
        content = section.get("content", "")

        # 章节标题
        heading = doc.add_heading(f"{i + 1}. {sec_title}", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        if sec_type == "table" and section.get("data"):
            # ── 表格 ──
            table_data = section["data"]
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])

            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # 表头
                for j, h in enumerate(headers):
                    cell = table.cell(0, j)
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

                # 数据行
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row):
                        if c_idx < len(headers):
                            cell = table.cell(r_idx + 1, c_idx)
                            cell.text = str(val)
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(10)

            # 表格后附加说明文字
            if content:
                doc.add_paragraph("")
                doc.add_paragraph(content)

        elif sec_type == "bullet" or "\n- " in content:
            # ── 列表 ──
            lines = content.split("\n")
            for line in lines:
                stripped = line.strip()
                if stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)

        else:
            # ── 普通文本 ──
            paragraphs = content.split("\n")
            for para_text in paragraphs:
                if para_text.strip():
                    # 检查是否是子标题（以 ### 开头）
                    if para_text.strip().startswith("### "):
                        heading = doc.add_heading(para_text.strip().lstrip("#").strip(), level=3)
                    elif para_text.strip().startswith("## "):
                        heading = doc.add_heading(para_text.strip().lstrip("#").strip(), level=2)
                    else:
                        doc.add_paragraph(para_text.strip())

    # ── 页脚（页码） ──
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加页码域
    from docx.oxml import OxmlElement
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run = footer_para.add_run()
    run._r.append(fldChar1)
    run2 = footer_para.add_run()
    run2._r.append(instrText)
    run3 = footer_para.add_run()
    run3._r.append(fldChar2)

    # 保存
    doc.save(output_path)
    print(f"✅ Word 文档已生成：{output_path}")
    print(f"   共 {len(sections)} 个章节")


def main():
    parser = argparse.ArgumentParser(description="Word 调研报告生成器")
    parser.add_argument("--input", required=True, help="内容 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    args = parser.parse_args()

    check_deps()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    output_dir = os.path.dirname(args.output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    create_docx(data, args.output)


if __name__ == "__main__":
    main()
